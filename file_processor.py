import PyPDF2
import docx
from PIL import Image
import io
import logging
from typing import Optional
import aiofiles
import os

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handle different file types and extract text content."""
    
    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                content = await file.read()
                
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            logger.info(f"Successfully extracted text from PDF: {len(text)} characters")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return None
    
    @staticmethod
    async def extract_text_from_docx(file_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {len(text)} characters")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return None
    
    @staticmethod
    async def extract_text_from_txt(file_path: str) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
            
            logger.info(f"Successfully extracted text from TXT: {len(text)} characters")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            return None
    
    @staticmethod
    async def process_image_with_ocr(file_path: str) -> Optional[str]:
        """Process image file and extract text using OCR (placeholder for future implementation)."""
        try:
            # For now, return a message indicating OCR is not implemented
            # In a full implementation, you would use libraries like pytesseract
            logger.warning("OCR processing not implemented yet")
            return "OCR text extraction from images is not yet implemented. Please upload PDF, DOCX, or TXT files."
            
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            return None
    
    @staticmethod
    async def process_file(file_path: str, file_extension: str) -> Optional[str]:
        """Process file based on its extension and extract text."""
        try:
            file_extension = file_extension.lower()
            
            if file_extension == '.pdf':
                return await FileProcessor.extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                return await FileProcessor.extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                return await FileProcessor.extract_text_from_txt(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                return await FileProcessor.process_image_with_ocr(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            return None
    
    @staticmethod
    def get_supported_extensions():
        """Return list of supported file extensions."""
        return ['.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.bmp', '.tiff']
    
    @staticmethod
    def is_supported_file(filename: str) -> bool:
        """Check if file extension is supported."""
        _, ext = os.path.splitext(filename.lower())
        return ext in FileProcessor.get_supported_extensions()
